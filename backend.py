from langchain_huggingface import HuggingFaceEndpoint,ChatHuggingFace
from langgraph.graph import StateGraph,START,END
from langchain_core.prompts import PromptTemplate
from pydantic import Field,BaseModel
import os
from typing import TypedDict,Annotated,List,Optional
from langchain_core.output_parsers import PydanticOutputParser

hf_token = os.getenv("HUGGINGFACEHUB_API_TOKEN")
llm=HuggingFaceEndpoint(repo_id="Qwen/Qwen2.5-7B-Instruct",huggingfacehub_api_token=hf_token)

model=ChatHuggingFace(llm=llm)

class Section(BaseModel):
    heading:str
    content:List[str]

class doc_schema(BaseModel):
    title:str=Field(description="the title of the documentation")
    section: List[Section] = Field(description="List of documentation sections")
class state(TypedDict):
    topic:Optional[str]
    context:Optional[str]
    all_headings:List[str]
    tone:str
    result:doc_schema
    file_path:str
    status:str

    

class decode_heading(BaseModel):
    headings:List[str]=Field(description="list of all the heading for the given content")
parser_decode=PydanticOutputParser(pydantic_object=decode_heading)
def decoder_node(state):
    if state.get("topic"):
        content=state["topic"]
    else:
        content=state["context"]
    tone=state["tone"]
    prompt=PromptTemplate(
        template="""
You are a senior technical architect and professional documentation specialist.

Your task is to analyze the provided project input and design the most appropriate
top-level documentation structure for it.

The project input may be:
- a short project topic, OR
- a detailed project description or context

You must dynamically determine which documentation sections are necessary
based on:
- the project domain and purpose
- system complexity
- technical depth
- intended audience

The tone of the documentation must strictly follow this user-defined tone:
{tone}

Guidelines:
- Do NOT use predefined or hardcoded documentation templates
- Do NOT reuse generic headings unless they are genuinely relevant
- Headings must be professional, meaningful, and specific to the project
- Avoid redundant or overlapping sections
- Decide the number of sections dynamically (typically 8–15 depending on complexity)
- Generate ONLY top-level section headings
- Do NOT generate explanations, descriptions, or subheadings

CRITICAL OUTPUT RULES:
- The response MUST strictly follow the provided output schema
- Return ONLY valid JSON
- Do NOT include markdown, numbering, or extra text
- Do NOT include anything outside the JSON object

Project Input:
{project_input}
\n
{format_instruction}
"""
,input_variables=["tone","project_input"],partial_variables={"format_instruction":parser_decode.get_format_instructions()}
    )    
    chain=prompt|model|parser_decode
    response=chain.invoke({"project_input":content,"tone":tone})

    return {"all_headings":response.headings}




parser=PydanticOutputParser(pydantic_object=doc_schema)

def topic_based(state):
    heading=state["all_headings"]
    final_heading = "\n".join(heading)
    topic=state["topic"]
    tone=state["tone"]
    prompt=PromptTemplate(
        template="""You are an expert technical documentation writer.

Your task is to generate COMPLETE, PROFESSIONAL project documentation
in the following tone: {tone}

The documentation must be detailed, comprehensive, and suitable for
developers and technical stakeholders.

========================
STRICT OUTPUT CONSTRAINTS (CRITICAL)
========================
- You MUST return ONLY valid JSON
- The JSON MUST strictly follow the provided schema
- Do NOT add extra keys, nesting, or structural changes
- Do NOT include markdown, comments, explanations, or formatting
- Do NOT include text outside the JSON object
- Every value must match the expected data type exactly

========================
STRUCTURE RULES (VERY IMPORTANT)
========================
- The title MUST be a single string
- Each section MUST contain:
  - "heading": a single string
  - "content": a list of strings ONLY
- Each content item must be a full, meaningful bullet-style explanation
- Do NOT create sub-sections, nested objects, or numbered keys
- Do NOT embed headings inside content strings

========================
CONTENT GUIDELINES
========================
- Generate multiple sections (minimum 8–10 or as appropriate)
- Each section must be relevant and non-redundant
- Each section must contain at least 4–6 detailed content points
- Expand concepts using best practices and general industry knowledge
- Keep language professional, technical, and practical
- If the topic is small, intelligently expand using common patterns

Use the following section headings as guidance.
You may include all of them and may add more if relevant,
but the structure must remain unchanged:

{final_heading}

========================
PROJECT TOPIC
========================
{topic}

========================
FINAL INSTRUCTION
========================
Generate the documentation strictly according to the schema below.
Any deviation from the schema will result in failure.

{format_instruction}

""",input_variables=["topic","tone","final_heading"],partial_variables={
    "format_instruction":parser.get_format_instructions()
} )
    
    chain=prompt|model|parser
    response=chain.invoke({"topic":topic,"tone":tone,"final_heading":final_heading})
    return {"result":response}


def context_based(state):
    context=state["context"]
    tone=state["tone"]
    heading=state["all_headings"]
    final_heading = "\n".join(heading)
    
    prompt=PromptTemplate(

        template="""
    You are an expert technical documentation architect.

Your task is to convert the given project context into a clean,
well-structured, and professional project documentation
using the following tone: {tone}

The input context may be unstructured, informal, incomplete, or brief.

If the context is:
- Detailed → reorganize and professionally rewrite it
- Brief or insufficient → expand it using realistic,
  industry-standard assumptions relevant to the topic

========================
STRICT OUTPUT CONSTRAINTS (CRITICAL)
========================
- You MUST return ONLY valid JSON
- The JSON MUST strictly follow the provided schema
- Do NOT add extra keys, nesting, or structural changes
- Do NOT include markdown, comments, or explanations
- Do NOT include text outside the JSON object

========================
STRUCTURE RULES (NON-NEGOTIABLE)
========================
- The title MUST be a single string
- Each section MUST contain ONLY:
  - "heading": string
  - "content": list of strings
- Each content item must be a complete, self-contained explanation
- Do NOT create sub-sections, nested objects, or grouped bullets
- Do NOT embed headings inside content strings
- Do NOT mirror the structure of the input context

========================
CONTENT GUIDELINES
========================
- Generate multiple logical sections (minimum 8–10 or as appropriate)
- Each section must contain at least 4–6 detailed bullet-style points
- Rewrite all content in a professional, technical tone
- Expand minimally provided context using common industry practices
- Avoid redundancy across sections
- Do NOT introduce speculative or advanced features

If the provided context is minimal, you may include standard sections such as:
{final_heading}

========================
PROJECT CONTEXT
========================
{context}

========================
FINAL INSTRUCTION
========================
Generate the documentation strictly according to the schema below.
Structural creativity is NOT allowed. Content expansion IS allowed.

{format_instruction}
 """,input_variables=["context","tone","final_heading"],partial_variables={"format_instruction":parser.get_format_instructions()}
)
    chain=prompt|model|parser
    response=chain.invoke({"context":context,"tone":tone,"final_heading":final_heading})
    return {"result":response}




from docx import Document
from datetime import datetime
import os
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def docx_generator(state):
    output=state["result"]
    doc=Document()

    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run(output.title)
    run.bold=True
    run.underline=True
    run.font.size=Pt(24)
    run.font.color.rgb=RGBColor(0, 51, 102)
    doc.add_paragraph()

    for section in output.section:
        doc.add_paragraph()
        s=doc.add_paragraph(style="List Number")
        run=s.add_run(section.heading)
        run.bold=True
        run.underline=True
        run.font.size=Pt(20)
        run.font.color.rgb=RGBColor(0, 0, 128)
        s.paragraph_format.space_after = Pt(6)

        
        for content in section.content:
            a=doc.add_paragraph(style="List Bullet")
            run=a.add_run(content)
            run.font.size=Pt(16)
            run.font.color.rgb=RGBColor(30,30,30)
            a.paragraph_format.space_before = Pt(4)
            a.paragraph_format.space_after = Pt(4)

                    
            

    filename = f"{output.title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    path=os.path.join("generated_docx",filename)
    os.makedirs("generated_docx", exist_ok=True)
    doc.save(path)
    return{"status":"done","file_path":path}
    


graph=StateGraph(state)
def condition(state):
    if(state.get("topic")):
        return "topic"
    else: return "context"


graph.add_node("topic",topic_based)
graph.add_node("doc",docx_generator)
graph.add_node("context",context_based)
graph.add_node("decode",decoder_node)
graph.add_edge(START,"decode")
graph.add_conditional_edges("decode", condition,({
    "topic":"topic","context":"context"
}))
graph.add_edge("topic","doc")
graph.add_edge("context","doc")
graph.add_edge("doc",END)

workflow=graph.compile()

